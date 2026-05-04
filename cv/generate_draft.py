"""Generate ATS-optimized Word CV draft for Santiago Aldana."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "output", "Santiago_Aldana_CV_ATS.docx")
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

doc = Document()

# ── Page margins (narrow for density) ──────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)

# ── Style helpers ───────────────────────────────────────────────────────────
def set_font(run, size, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_space(p, before=0, after=0):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)

def add_section_header(doc, text):
    p = doc.add_paragraph()
    para_space(p, before=8, after=2)
    run = p.add_run(text.upper())
    set_font(run, 9.5, bold=True, color=(26, 26, 46))
    # Bottom border (thin rule under section title)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1a1a2e')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_bullet(doc, text, indent=0.18):
    p = doc.add_paragraph(style='List Bullet')
    para_space(p, before=1, after=1)
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    run = p.add_run(text)
    set_font(run, 9)
    return p

def add_achievement(doc, text):
    p = doc.add_paragraph()
    para_space(p, before=2, after=2)
    p.paragraph_format.left_indent = Inches(0.18)
    run_label = p.add_run("Key Achievement: ")
    set_font(run_label, 9, bold=True)
    run_text = p.add_run(text)
    set_font(run_text, 9)
    return p

def add_role_header(doc, title, company, location, dates):
    # Title on its own line — ATS parsers need title and company separated
    p1 = doc.add_paragraph()
    para_space(p1, before=6, after=0)
    r1 = p1.add_run(title)
    set_font(r1, 9.5, bold=True)
    # Company | Location | Dates on the next line
    p2 = doc.add_paragraph()
    para_space(p2, before=0, after=1)
    r2 = p2.add_run(f"{company}  |  {location}  |  {dates}")
    set_font(r2, 9, color=(100, 100, 100))
    return p2

# ════════════════════════════════════════════════════════════════════════════
# HEADER
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(p, after=2)
r = p.add_run("Santiago Aldana")
set_font(r, 18, bold=True, color=(26, 26, 46))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(p, after=6)
r = p.add_run("617-216-7003  |  aldana.santiago@gmail.com  |  linkedin.com/in/santiago-aldana  |  Boston, MA")
set_font(r, 9, color=(60, 60, 60))

# ════════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════
add_section_header(doc, "Executive Summary")

p = doc.add_paragraph()
para_space(p, before=3, after=3)
r = p.add_run(
    "MIT Sloan MBA executive with 20+ years of strategic leadership in fintech, AI, payments, and digital "
    "transformation across Latin America and the U.S. Founder and operator with two successful exits; proven "
    "record building high-velocity platforms inside both startups and regulated enterprises. Currently applying "
    "agentic AI, embedded finance, and BaaS expertise at the intersection of fintech innovation and institutional scale."
)
set_font(r, 9)

# ════════════════════════════════════════════════════════════════════════════
# CORE COMPETENCIES  (ATS keyword block — individual terms, not phrases)
# ════════════════════════════════════════════════════════════════════════════
add_section_header(doc, "Core Competencies")

competencies = [
    "Product Strategy", "Payments Infrastructure", "Fraud Prevention", "Digital Identity",
    "KYC / AML Compliance", "Biometrics & Authentication", "Open Banking", "Embedded Finance",
    "Banking-as-a-Service (BaaS)", "Agentic AI", "AI/ML Strategy", "Blockchain & Stablecoins",
    "Cross-border Payments", "Financial Inclusion", "Go-to-Market Strategy", "Strategic Partnerships",
    "P&L Ownership", "Digital Transformation", "Agile / Lean Startup", "Risk Management",
]
p = doc.add_paragraph()
para_space(p, before=3, after=3)
r = p.add_run("  •  ".join(competencies))
set_font(r, 8.5)

# ════════════════════════════════════════════════════════════════════════════
# PROFESSIONAL EXPERIENCE
# ════════════════════════════════════════════════════════════════════════════
add_section_header(doc, "Professional Experience")

# ── SMCU ─────────────────────────────────────────────────────────────────
add_role_header(doc,
    "Chief Product & Solutions Officer",
    "St. Mary's Credit Union",
    "Marlborough, MA",
    "January 2026 – Present"
)
bullets = [
    "Architecting a fintech partnership framework and CUSO strategy to embed AI-native financial services — including agentic AI for CX, fraud prevention, and marketing intelligence — into a federally regulated credit union.",
    "Evaluating and deploying embedded finance and BaaS solutions, bridging the speed of fintech innovation with the compliance and governance requirements of a regulated institution.",
    "Building the product and go-to-market strategy for new revenue lines at the intersection of payments, digital identity, and AI-driven member experience.",
    "Applying 20+ years of fintech operator experience to identify and execute strategic partnerships that accelerate capabilities traditional institutions would take years to build internally.",
]
for b in bullets:
    add_bullet(doc, b)
add_achievement(doc,
    "Leading SMCU's first structured fintech partnership program, positioning the credit union as a platform for embedded finance innovation in the Greater Boston market."
)

# ── Strategic Advisor ────────────────────────────────────────────────────
add_role_header(doc,
    "Strategic Advisor — Fintech & Emerging Tech",
    "Independent",
    "Boston, MA & Bogotá, Colombia (Remote)",
    "March 2022 – December 2025"
)
bullets = [
    "Provided board-level guidance and technical strategy for growth-stage fintechs and financial institutions, focused on Banking-as-a-Service (BaaS), Agentic AI, and cross-border payment infrastructure.",
    "TUYA (Board Member): Directed the transition from a traditional retail lender to a BaaS provider; leveraged retail transaction data to develop digital products for underserved market segments.",
    "MAVEN AGI: Led LATAM market exploration, building the distribution channel for AI agent deployment to automate customer support and localize go-to-market operations.",
    "ZULU (Board Member): Developed cross-border payment rails using stablecoins to reduce settlement times and FX costs between the U.S. and Latin America.",
    "AVANZO: Guided development of inclusive lending platforms using alternative data sources to evaluate and price credit risk for underserved populations.",
    "COLOMBIA FINTECH (Board Member): Collaborated with the Central Bank and regulators to define standards for Open Finance and Instant Payment regulations (Bre-B).",
]
for b in bullets:
    add_bullet(doc, b)
add_achievement(doc,
    "Concurrent advisory portfolio spanning BaaS, Agentic AI, stablecoin payments, and Open Finance regulation — providing hands-on strategic guidance across 5 organizations simultaneously."
)

# ── SoyYo ────────────────────────────────────────────────────────────────
add_role_header(doc,
    "CEO",
    "SoyYo — Digital Identity & Fraud Prevention Platform",
    "Bogotá, Colombia",
    "2020 – 2024"
)
bullets = [
    "Created and scaled a user-centric digital identity platform adopted by 3M+ users, delivering onboarding, KYC, biometrics, authentication, and cybersecurity services; negotiated its successful sale to Redeban, Colombia's leading payment processor.",
    "Forged strategic alliances with Mastercard, LexisNexis, Experian, Thales, and TransUnion, positioning SoyYo as a trusted ecosystem for digital identity and fraud prevention.",
    "Directed cross-bank collaboration among Colombia's top 3 financial institutions, aligning stakeholders on AI-driven biometrics, authentication, consent management, and fraud prevention in highly regulated financial products.",
    "Built a high-performing digital team from scratch, achieving an employee NPS of 9.1 and fostering a culture of innovation and execution.",
]
for b in bullets:
    add_bullet(doc, b)
add_achievement(doc,
    "Founded and scaled SoyYo's API-first identity platform from zero to 3M+ user identities; led successful acquisition by Colombia's leading PSP (Redeban)."
)

# ── Avianca ──────────────────────────────────────────────────────────────
add_role_header(doc,
    "Chief Digital & Technology Officer",
    "Avianca",
    "Bogotá, Colombia",
    "2017 – 2019"
)
bullets = [
    "Directed a $110M IT budget, delivering a company-wide digital transformation that enabled omnichannel customer knowledge and personalization.",
    "Grew online sales by 15% YoY, generating $700–800M annually and representing 47% of Avianca's total sales growth.",
    "Enabled Amazon partnership, positioning Avianca Cargo as Latin America's primary e-commerce distribution partner.",
    "Improved digital NPS by 12+ points by redesigning customer touchpoints across all digital channels.",
    "Embedded analytics, data, and innovation into core operations through design thinking, agile cells, and lean startup methodologies.",
]
for b in bullets:
    add_bullet(doc, b)
add_achievement(doc,
    "Delivered airline-wide digital transformation — 47% of total sales growth attributed to digital channels, NPS up 12 points, payments conversion up 10 percentage points."
)

# ── IQ Outsourcing ───────────────────────────────────────────────────────
add_role_header(doc,
    "CEO",
    "IQ Outsourcing — BPO & IT Services (Banking, Telecom, Insurance)",
    "Bogotá, Colombia",
    "2015 – 2017"
)
bullets = [
    "Transformed the business from a traditional BPO into a digital solutions provider for banking, health, and telecom sectors.",
    "Reduced HR cost-to-revenue ratio by 21 percentage points through AI, automation, and IT-driven process efficiency.",
    "Launched new digital products for banking and health industries, enhancing customer knowledge and engagement.",
    "Pioneered adoption of Blockchain, AI, mobile, DevOps, and cloud technologies across the client portfolio.",
]
for b in bullets:
    add_bullet(doc, b)
add_achievement(doc,
    "Improved credit card conversion rate at retail through AI-driven digital transformation of the issuing and onboarding process."
)

# ── Uff Móvil ────────────────────────────────────────────────────────────
add_role_header(doc,
    "CEO & Founding Partner",
    "Uff Móvil",
    "Bogotá, Colombia",
    "2010 – 2015"
)
bullets = [
    "Founded and launched Latin America's first Mobile Virtual Network Operator (MVNO), expanding financial inclusion through mobile tools, education, and digital channels.",
    "Scaled to 400K customers in the first year, demonstrating rapid product-market fit and execution.",
    "Used long-distance call patterns and mobile top-up behavior as proxies for remittance intent and credit risk — pioneering behavioral data scoring before the category existed.",
    "Negotiated the sale of 70% of the company to Bancolombia at an $18M valuation; subsequently advised Bancolombia Group ventures on mobile, digital, and fintech innovation.",
]
for b in bullets:
    add_bullet(doc, b)
add_achievement(doc,
    "Pioneering Fintech — Founded Latin America's first MVNO; grew to 400K customers in year one, valued at $18M, acquired by Bancolombia."
)

# ── Telefónica ───────────────────────────────────────────────────────────
add_role_header(doc,
    "CIO & Director of Customer Loyalty",
    "Telefónica",
    "Bogotá, Colombia",
    "2004 – 2009"
)
bullets = [
    "Delivered a €60M IT transformation program across 5 countries in 17 months (against a 4-year plan), modernizing operations and enabling rapid product launches.",
    "Enabled 100%+ annual sales growth by improving provisioning systems and expanding product offerings.",
    "Led the operational and technological merger of 40 business units, positioning the company for acquisition by Telefónica.",
    "Directed 500+ professionals across multiple countries, improving customer retention and sales conversion by 5+ percentage points through broadband and TV strategies.",
]
for b in bullets:
    add_bullet(doc, b)
add_achievement(doc,
    "Technology at Scale — Led Telefónica's €60M IT transformation across 5 countries in 17 months (vs. a 4-year plan)."
)

# ════════════════════════════════════════════════════════════════════════════
# PREVIOUS EXPERIENCE
# ════════════════════════════════════════════════════════════════════════════
add_section_header(doc, "Earlier Experience")

add_bullet(doc,
    "Chamber of Commerce of Bogotá — CFO: Led digital transformation through SAP implementation; "
    "managed a digital certificate subsidiary as Board Member; introduced risk-based model for a COP 70B investment portfolio."
)
add_bullet(doc,
    "Citibank — Marketing & Sales Manager: Launched Mastercard and Visa credit cards in Colombia; "
    "reduced onboarding time from two weeks to one day; developed the 'Citigold' high-value customer segment."
)

# ════════════════════════════════════════════════════════════════════════════
# BOARD & ADVISORY  (consolidated — detail now in Strategic Advisor role)
# ════════════════════════════════════════════════════════════════════════════
add_section_header(doc, "Board & Advisory Roles")

board_items = [
    ("Zulu", "Board Member", "2024–2025", "Cross-border payments platform using crypto and stablecoins."),
    ("Tuya Credit Card", "Independent Board Member", "2021–2024", "Colombia's largest card issuer (bank + retailer JV); guided Open Banking and embedded finance strategy."),
    ("Colombia Fintech", "Board Member", "2019–2022", "Led regulatory advocacy for Open Finance, instant payments, and financial inclusion."),
]
for company, role, dates, summary in board_items:
    p = doc.add_paragraph()
    para_space(p, before=2, after=1)
    p.paragraph_format.left_indent = Inches(0.18)
    r1 = p.add_run(f"{company} — {role}  ({dates}): ")
    set_font(r1, 9, bold=True)
    r2 = p.add_run(summary)
    set_font(r2, 9)

# ════════════════════════════════════════════════════════════════════════════
# EDUCATION
# ════════════════════════════════════════════════════════════════════════════
add_section_header(doc, "Education")

edu = [
    ("MIT Sloan School of Management", "Boston, MA", "MBA — Management of Technology, Strategy, Innovation & Technology"),
    ("Universidad de los Andes", "Bogotá, Colombia", "Bachelor's in Industrial Engineering — Thesis: Credit Score Model for Diners Club Credit Card"),
]
for inst, loc, degree in edu:
    p = doc.add_paragraph()
    para_space(p, before=2, after=1)
    p.paragraph_format.left_indent = Inches(0.18)
    r1 = p.add_run(f"{inst}, {loc}  — ")
    set_font(r1, 9, bold=True)
    r2 = p.add_run(degree)
    set_font(r2, 9)

# ════════════════════════════════════════════════════════════════════════════
# ADDITIONAL
# ════════════════════════════════════════════════════════════════════════════
add_section_header(doc, "Additional Information")
p = doc.add_paragraph()
para_space(p, before=2)
p.paragraph_format.left_indent = Inches(0.18)
r = p.add_run("Languages: Fully bilingual in Spanish and English  |  Work Authorization: U.S.A. and Colombia")
set_font(r, 9)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
